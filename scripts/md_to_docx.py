#!/usr/bin/env python3
"""Convert markdown files to styled DOCX suitable for Google Drive upload."""

import markdown
import sys
import base64
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree


def _set_cell_shading(cell, color_hex):
    """Apply shading to a table cell."""
    shading = etree.SubElement(
        cell._element.get_or_add_tcPr(),
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd'
    )
    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', color_hex)
    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')


def _set_paragraph_shading(paragraph, color_hex):
    """Apply shading to a paragraph background."""
    pPr = paragraph._element.get_or_add_pPr()
    shading = etree.SubElement(
        pPr,
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd'
    )
    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', color_hex)
    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')


def _add_bottom_border(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
    bottom = etree.SubElement(pBdr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '4')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', 'DADCE0')


def convert(md_path):
    with open(md_path, 'r') as f:
        md_content = f.read()

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Arial'
        if i == 1:
            hs.font.size = Pt(18)
            hs.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        elif i == 2:
            hs.font.size = Pt(14)
            hs.font.color.rgb = RGBColor(0x20, 0x21, 0x24)
        elif i == 3:
            hs.font.size = Pt(12)
            hs.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        if line.startswith('```') and not in_code_block:
            if in_table:
                _flush_table(doc, table_rows)
                in_table = False
                table_rows = []
            in_code_block = True
            code_lines = []
            i += 1
            continue

        if line.startswith('```') and in_code_block:
            in_code_block = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
            _set_paragraph_shading(p, 'F1F3F4')
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if '|' in line and not line.strip().startswith('```'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue

        if in_table:
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []

        if line.startswith('═') or line.strip() == '':
            i += 1
            continue

        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            _add_bottom_border(p)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            _add_bottom_border(p)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('---'):
            p = doc.add_paragraph()
            _add_bottom_border(p)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_formatting(p, line[2:])
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            _add_inline_formatting(p, re.sub(r'^\d+\.\s', '', line))
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _set_paragraph_shading(p, 'E8F0FE')
            _add_inline_formatting(p, line[2:])
        elif line.strip():
            p = doc.add_paragraph()
            _add_inline_formatting(p, line)

        i += 1

    if in_table:
        _flush_table(doc, table_rows)

    out_path = md_path.rsplit('.', 1)[0] + '.docx'
    doc.save(out_path)

    with open(out_path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8'), out_path


def _flush_table(doc, rows):
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for ri, row_data in enumerate(rows):
        row = table.rows[ri]
        for ci, cell_text in enumerate(row_data):
            if ci < num_cols:
                cell = row.cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(cell_text)
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_cell_shading(cell, '1A73E8')
                elif ri % 2 == 0:
                    _set_cell_shading(cell, 'F8F9FA')


def _add_inline_formatting(paragraph, text):
    paragraph.clear()
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`|→)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
        elif part == '→':
            run = paragraph.add_run('→')
            run.font.name = 'Arial'
            run.font.size = Pt(11)
        elif part:
            run = paragraph.add_run(part)
            run.font.name = 'Arial'
            run.font.size = Pt(11)


if __name__ == '__main__':
    b64, path = convert(sys.argv[1])
    print(f"Saved: {path}")
    print(b64)
